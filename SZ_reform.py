import sys

import customtkinter as CTk
from docx import Document
from docx.shared import Pt
import CTkMessagebox
import os.path


def on_document_sz_select(root, second, fio, phn):
    second.withdraw()
    CTk.set_default_color_theme("blue")

    third = CTk.CTk()
    third.geometry(
        "400x750+{}+{}".format(int(third.winfo_screenwidth() / 2 - 150), int(third.winfo_screenheight() / 2 - 300)))
    third.resizable(width=False, height=False)
    third.title("Служебная записка")
    third.iconbitmap('капля.ico')

    address_label = CTk.CTkLabel(third, text="Выберите адресата:")
    address_label.pack()

    address_var = CTk.StringVar()
    address_combobox = CTk.CTkComboBox(third, width=300, variable=address_var, corner_radius=4, border_width=1,
                                       border_color="#1f6aa5", button_hover_color='#1f6aa5',
                                       dropdown_hover_color='#1f6aa5',
                                       values=["Генеральный директор",
                                               "Главный инженер",
                                               "Директор по безопасности",
                                               "Директор по персоналу",
                                               "Начальник технического отдела",
                                               "Начальник отдела оптимизации режимов сетей и сооружений",
                                               "Начальник службы АСУ ТП и Метрологии",
                                               "Начальник производственного отдела",
                                               "Заместитель генерального директора Производственной дирекции",
                                               "Начальник центральной диспетчерской службы",
                                               "Заместитель генерального директора Дирекции по капитальному строительству",
                                               "Коммерческий директор",
                                               "Директор по экономике",
                                               "Заместитель генерального директора Дирекции по правовым вопросам",
                                               "Начальник отдела контроллинга",
                                               "Начальник отдела информационных технологий",
                                               "Начальник управления системой снабжения"])
    address_combobox.pack()

    theme_label = CTk.CTkLabel(third, text="Введите тему:")
    theme_label.pack()

    theme_entry = CTk.CTkEntry(third, width=300, corner_radius=4, border_width=1,
                               border_color="#1f6aa5")
    theme_entry.pack()

    content_label = CTk.CTkLabel(third, text="Введите содержание:")
    content_label.pack()

    def on_paste():
        content_entry.paste()

    content_entry = CTk.CTkTextbox(third, height=200, width=300, corner_radius=4, border_width=1,
                                   border_color="#1f6aa5", fg_color='gray19')
    content_entry.bind('<Control-v>', on_paste)
    content_entry.pack()

    save_path_label = CTk.CTkLabel(third, text="Выберите место сохранения:")
    save_path_label.pack()

    save_path_entry = CTk.CTkEntry(third, width=300, corner_radius=4, border_width=1,
                                   border_color="#1f6aa5", state='disabled')
    save_path_entry.pack()
    save_path_entry.place(x=50, y=370)

    def select_save_path():
        save_path_entry.configure(state='normal')
        save_path = CTk.filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")])
        save_path_entry.delete(0, "end")
        save_path_entry.insert(0, save_path)
        save_path_entry.configure(state='disabled')
        generate_button.configure(state='normal')

    select_button = CTk.CTkButton(third, text="Выбрать место сохранения", corner_radius=10, command=select_save_path)
    select_button.pack()
    select_button.place(x=110, y=420)

    def sz_reform_def():
        theme = theme_entry.get()
        choice = address_combobox.get()
        content = content_entry.get("1.0", "end-1c")
        save_path = save_path_entry.get()
        doc = Document('СЗ_шаблон.docx')
        section = doc.sections[-1]
        footer = section.footer
        footer_para = footer.paragraphs[-1]
        footer_para.text = "Исп. " + fio.get() + ' ' + 'тел. ' + phn.get()
        run = footer_para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        doc.save('Temp.docx')

        docu = Document('Temp.docx')

        for table in docu.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "(Адресат)" in paragraph.text:
                            if choice == "Генеральный директор":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Генеральному директору")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Главный инженер":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Главному инженеру")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по безопасности":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Директору по безопасности")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по персоналу":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Директору по персоналу")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник технического отдела":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику технического отдела")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Главный технолог":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Главному технологу")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела оптимизации режимов сетей и сооружений":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику отдела оптимизации режимов сетей и сооружений")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник службы АСУ ТП и Метрологии":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику службы АСУ ТП и Метрологии")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник производственного отдела":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику производственного отдела")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Производственной дирекции":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Заместителю генерального директора Производственной дирекции")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник центральной диспетчерской службы":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику центральной диспетчерской службы")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Дирекции по капитальному строительству":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Заместителю генерального директора Дирекции по капитальному строительству")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Коммерческий директор":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Коммерческому директору")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по экономике":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Директору по экономике")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Дирекции по правовым вопросам":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Заместителю генерального директора Дирекции по правовым вопросам")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела контроллинга":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику отдела контроллинга")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела информационных технологий":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику отдела информационных технологий")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник управления системой снабжения":
                                paragraph.text = paragraph.text.replace("(Адресат)", "Начальнику управления системой снабжения")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)

                        if "(Имя и Отчество)" in paragraph.text:
                            if choice == "Генеральный директор":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Павел Гаврилович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Главный инженер":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Денис Александрович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по безопасности":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Андрей Владимирович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по персоналу":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Анна Анатольевна")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник технического отдела":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Сергей Анатольевич")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Главный технолог":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Евгений Борисович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела оптимизации режимов сетей и сооружений":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Денис николаевич")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник службы АСУ ТП и Метрологии":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Александр Олегович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник производственного отдела":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Ирина Александровна")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Производственной дирекции":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Сергей Владимирович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник центральной диспетчерской службы":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Дмитрий Григорьевич")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Дирекции по капитальному строительству":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Евгений Гарьевич")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Коммерческий директор":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Ольга Юрьевна")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Директор по экономике":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Михаил Александрович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Заместитель генерального директора Дирекции по правовым вопросам":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Дмитрий Анатольевич")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела контроллинга":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Оксана Романовна")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник отдела информационных технологий":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Сергей Вадимович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)
                            elif choice == "Начальник управления системой снабжения":
                                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Дмитрий Александрович")
                                run = paragraph.runs[0]
                                run.font.name = "Arial"
                                run.font.size = Pt(14)

                        if "(Тема)" in paragraph.text:
                            paragraph.text = paragraph.text.replace("(Тема)", theme)
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)

        # Добавим замену для текста, не находящегося в таблице
        for paragraph in docu.paragraphs:
            if choice == "Генеральный директор":
                paragraph.text = paragraph.text.replace("(Адресат)", "Генеральному директору")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            elif choice == "Главный инженер":
                paragraph.text = paragraph.text.replace("(Адресат)", "Главный инженер")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            elif choice == "Директор по безопасности":
                paragraph.text = paragraph.text.replace("(Адресат)", "Директор по безопасности")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            if "(Имя и Отчество)" in paragraph.text:
                if choice == "Генеральный директор":
                    paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Павел Гаврилович")
                    run = paragraph.runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(14)
                elif choice == "Главный инженер":
                    paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Денис Александрович")
                    run = paragraph.runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(14)
                elif choice == "Директор по безопасности":
                    paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Андрей Владимирович")
                    run = paragraph.runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(14)
                if "(Тема)" in paragraph.text:
                    paragraph.text = paragraph.text.replace("(Тема)", theme)
                    run = paragraph.runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(14)

            if "(Содержание)" in paragraph.text:
                paragraph.text = paragraph.text.replace("(Содержание)", content)
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)

        docu.save(save_path)
        os.remove("Temp.docx")
        msg = CTkMessagebox.CTkMessagebox(
            title="Docs Generator",
            message="Документ успешно сформирован! Хотите сформировать еще?",
            icon='check',
            option_1='Да',
            option_2='Нет')

        if msg.get() == 'Да':
            third.destroy()
            second.deiconify()
        if msg.get() == 'Нет':
            sys.exit()

    generate_button = CTk.CTkButton(third, text="Сформировать документ", command=sz_reform_def, state='disabled',
                                    width=200, height=70, fg_color='#c52222', hover_color='#9d1b1b', corner_radius=10)
    generate_button.pack()
    generate_button.place(x=105, y=550)

    def on_closing():
        sys.exit()

    def back():
        third.destroy()
        second.deiconify()

    arrow = CTk.CTkButton(third, width=25, height=25, text="←", command=back)
    arrow.pack()
    arrow.place(x=15, y=710)

    third.protocol("WM_DELETE_WINDOW", on_closing)
    third.mainloop()
