import sys

import customtkinter as CTk
from docx import Document
from docx.shared import Pt
import CTkMessagebox
import os.path


def on_document_prikaz_select(root, second, fio, phn):
    second.withdraw()
    CTk.set_default_color_theme("blue")

    third = CTk.CTk()
    third.geometry(
        "400x750+{}+{}".format(int(third.winfo_screenwidth() / 2 - 150), int(third.winfo_screenheight() / 2 - 300)))
    third.resizable(width=False, height=False)
    third.title("Приказ")
    third.iconbitmap('капля.ico')

    theme_label = CTk.CTkLabel(third, text="Введите тему:")
    theme_label.pack()

    theme_entry = CTk.CTkEntry(third, width=300, corner_radius=4, border_width=1,
                               border_color="#1f6aa5")
    theme_entry.pack()

    content_label = CTk.CTkLabel(third, text="Введите содержание:")
    content_label.pack()

    def on_paste():
        content_entry.paste()

    content_entry = CTk.CTkTextbox(third, height=200, width=300, corner_radius=4, border_width=1,
                                   border_color="#1f6aa5")
    content_entry.bind('<Control-v>', on_paste)
    content_entry.pack()

    save_path_label = CTk.CTkLabel(third, text="Выберите место сохранения:")
    save_path_label.pack()

    def select_save_path():
        save_path_entry.configure(state='normal')
        save_path = CTk.filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")])
        save_path_entry.delete(0, "end")
        save_path_entry.insert(0, save_path)
        save_path_entry.configure(state='disabled')
        generate_button.configure(state='normal')

    save_path_entry = CTk.CTkEntry(third, width=300, corner_radius=4, border_width=1,
                               border_color="#1f6aa5", state='disabled')
    save_path_entry.pack()
    save_path_entry.place(x=50, y=315)

    select_button = CTk.CTkButton(third, text="Выбрать место сохранения", corner_radius=10, command=select_save_path)
    select_button.pack()
    select_button.place(x=110, y=370)

    def sz_reform_def():
        theme = theme_entry.get()
        content = content_entry.get("1.0", "end-1c")
        save_path = save_path_entry.get()
        doc = Document('Приказ_шаблон.docx')
        section = doc.sections[-1]
        footer = section.footer
        footer_para = footer.paragraphs[-1]
        footer_para.text = "Исп. " + fio.get() + ' ' + 'тел. ' + phn.get()
        run = footer_para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        doc.save('Temp.docx')

        docu = Document('Temp.docx')

        for table in docu.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "(Тема)" in paragraph.text:
                            paragraph.text = paragraph.text.replace("(Тема)", theme)
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)

        # Добавим замену для текста, не находящегося в таблице
        for paragraph in docu.paragraphs:
            if "(Содержание)" in paragraph.text:
                paragraph.text = paragraph.text.replace("(Содержание)", content)
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)

        docu.save(save_path)
        os.remove("Temp.docx")
        msg = CTkMessagebox.CTkMessagebox(
            title="Docs Generator",
            message="Документ успешно сформирован! Хотите сформировать еще?",
            icon='check',
            option_1='Да',
            option_2='Нет')

        if msg.get() == 'Да':
            third.destroy()
            second.deiconify()
        if msg.get() == 'Нет':
            sys.exit()

    generate_button = CTk.CTkButton(third, text="Сформировать документ", command=sz_reform_def, state='disabled',
                                    width=200, height=70, fg_color='#c52222', hover_color='#9d1b1b', corner_radius=10)
    generate_button.pack()
    generate_button.place(x=105, y=550)

    def on_closing():
        sys.exit()

    def back():
        third.destroy()
        second.deiconify()

    arrow = CTk.CTkButton(third, width=25, height=25, text="←", command=back)
    arrow.pack()
    arrow.place(x=15, y=710)

    third.protocol("WM_DELETE_WINDOW", on_closing)
    third.mainloop()
