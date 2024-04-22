# Библиотеки
import os.path
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog

import docx
from docx import Document
from docx.shared import Pt
import customtkinter
from PIL import Image, ImageTk
import CTkMessagebox

# Загрузка картинок для кнопки
file_path = os.path.dirname(os.path.realpath((__file__)))
image_1 = customtkinter.CTkImage(Image.open(file_path + "/dark.png"), size=(35, 35))
image_2 = customtkinter.CTkImage(Image.open(file_path + "/light.png"), size=(35, 35))


def find_replace_in_table(table, find_text, replace_text):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if find_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(find_text, replace_text)


def on_document_type_select():
    if type_combobox.get() == "Служебная записка":
        print("Выбрана служебка")
        addr_entry.configure(state="disabled")
        answer_entry.configure(state="disabled")
        male_radio.configure(state="disabled")
        female_radio.configure(state="disabled")

    if type_combobox.get() == "Приказ":
        print("Выбрана служебка")
        addr_entry.configure(state="disabled")
        answer_entry.configure(state="disabled")
        address_combobox.configure(state="disabled")
        male_radio.configure(state="disabled")
        female_radio.configure(state="disabled")
        content_entry.configure(state="disabled")

    if type_combobox.get() == "Письмо":
        addr_entry.configure(state="normal")
        answer_entry.configure(state="normal")
        address_combobox.configure(state="normal")
        male_radio.configure(state="normal")
        female_radio.configure(state="normal")
        content_entry.configure(state="normal")


def generate_document():
    address = address_var.get()
    theme = theme_entry.get()
    content = content_entry.get("1.0", "end-1c")
    answer = answer_entry.get()
    addr = addr_entry.get()
    save_path = save_path_entry.get()

    doc = Document("Temp.docx")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "(Адресат)" in paragraph.text:
                        if address == "Генеральный директор":
                            paragraph.text = paragraph.text.replace("(Адресат)", "Генеральный директор")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)
                        elif address == "Главный инженер":
                            paragraph.text = paragraph.text.replace("(Адресат)", "Главный инженер")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)
                        elif address == "Директор по безопасности":
                            paragraph.text = paragraph.text.replace("(Адресат)", "Директор по безопасности")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)

                    if "(Имя и Отчество)" in paragraph.text:
                        if address == "Генеральный директор":
                            paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Павел Гаврилович")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)
                        elif address == "Главный инженер":
                            paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Денис Александрович")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)
                        elif address == "Директор по безопасности":
                            paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Андрей Владимирович")
                            run = paragraph.runs[0]
                            run.font.name = "Arial"
                            run.font.size = Pt(14)

                    if "(Тема)" in paragraph.text:
                        paragraph.text = paragraph.text.replace("(Тема)", theme)
                        run = paragraph.runs[0]
                        run.font.name = "Arial"
                        run.font.size = Pt(14)

                    if "(Содержание)" in paragraph.text:
                        paragraph.text = paragraph.text.replace("(Содержание)", content)
                        run = paragraph.runs[0]
                        run.font.name = "Arial"
                        run.font.size = Pt(14)

                    if "(На какое письмо)" in paragraph.text:
                        paragraph.text = paragraph.text.replace("(На какое письмо)", answer)
                        run = paragraph.runs[0]
                        run.font.name = "Arial"
                        run.font.size = Pt(14)

                    if "(Адрес)" in paragraph.text:
                        paragraph.text = paragraph.text.replace("(Адрес)", addr)
                        run = paragraph.runs[0]
                        run.font.name = "Arial"
                        run.font.size = Pt(14)

                    if gender_var.get() == "Мужской":
                        paragraph.text = paragraph.text.replace("(Пол)", "ый")
                    elif gender_var.get() == "Женский":
                        paragraph.text = paragraph.text.replace("(Пол)", "ая")

    # Добавим замену для текста, не находящегося в таблице
    for paragraph in doc.paragraphs:
        if "(Имя и Отчество)" in paragraph.text:
            if address == "Генеральный директор":
                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Павел Гаврилович")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            elif address == "Главный инженер":
                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Денис Александрович")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            elif address == "Директор по безопасности":
                paragraph.text = paragraph.text.replace("(Имя и Отчество)", "Андрей Владимирович")
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)
            if "(Тема)" in paragraph.text:
                paragraph.text = paragraph.text.replace("(Тема)", theme)
                run = paragraph.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(14)

        if "(Содержание)" in paragraph.text:
            paragraph.text = paragraph.text.replace("(Содержание)", content)
            run = paragraph.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(14)

        if gender_var.get() == "Мужской":
            paragraph.text = paragraph.text.replace("(Пол)", "ый")
            run = paragraph.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(14)
        elif gender_var.get() == "Женский":
            paragraph.text = paragraph.text.replace("(Пол)", "ая")
            run = paragraph.runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(14)

    doc.save(save_path)
    os.remove("Temp.docx")
    CTkMessagebox.CTkMessagebox(title="Docs Generator", message="Документ успешно сформирован!")
    address = None
    theme = None
    content = None
    answer = None
    addr = None
    save_path = None


# Создание графического интерфейса
customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")

root = customtkinter.CTk()
root.geometry("400x750")
root.resizable(width=False, height=False)
root.title("Генератор документов")

type_label = customtkinter.CTkLabel(root, text="Выберите вид документа:")
type_label.pack()

type_var = customtkinter.StringVar()
type_combobox = customtkinter.CTkComboBox(root, width=300, variable=type_var, values=["Служебная записка", "Приказ", "Письмо"])
type_combobox.pack()
type_combobox.bind("<<ComboboxSelected>>", lambda event: on_document_type_select())

address_label = customtkinter.CTkLabel(root, text="Выберите адресата:")
address_label.pack()

gender_var = tk.StringVar()
gender_var.set("Мужской")
male_radio = customtkinter.CTkRadioButton(root, text="Мужской", variable=gender_var, radiobutton_width=15,
                                          radiobutton_height=15, value="Мужской")
male_radio.pack()
female_radio = customtkinter.CTkRadioButton(root, text="Женский", variable=gender_var, radiobutton_width=15,
                                            radiobutton_height=15, value="Женский")
female_radio.pack()

address_var = tk.StringVar()
address_combobox = customtkinter.CTkComboBox(root, width=300, variable=address_var,
                                             values=["Генеральный директор", "Главный инженер",
                                                     "Директор по безопасности"])
address_combobox.pack()

theme_label = customtkinter.CTkLabel(root, text="Введите тему:")
theme_label.pack()

theme_entry = customtkinter.CTkEntry(root, width=300)
theme_entry.pack()

answer_label = customtkinter.CTkLabel(root, text="Введите номер и дату входящего письма:")
answer_label.pack()

answer_entry = customtkinter.CTkEntry(root, width=300)
answer_entry.pack()

addr_label = customtkinter.CTkLabel(root, text="Введите адрес получателя:")
addr_label.pack()

addr_entry = customtkinter.CTkEntry(root, width=300)
addr_entry.pack()

content_label = customtkinter.CTkLabel(root, text="Введите содержание:")
content_label.pack()


def on_paste():
    content_entry.paste()


content_entry = customtkinter.CTkTextbox(root, height=200, width=300)
content_entry.bind('<Control-v>', on_paste)
content_entry.pack()

save_path_label = customtkinter.CTkLabel(root, text="Выберите место сохранения:")
save_path_label.pack()

save_path_entry = customtkinter.CTkEntry(root, width=300)
save_path_entry.pack()


def select_save_path():
    save_path = customtkinter.filedialog.asksaveasfilename(defaultextension=".docx",
                                                           filetypes=[("Word files", "*.docx")])
    save_path_entry.delete(0, "end")
    save_path_entry.insert(0, save_path)


select_button = customtkinter.CTkButton(root, text="Выбрать место сохранения", command=select_save_path)
select_button.pack()
select_button.place(x=110, y=615)

generate_button = customtkinter.CTkButton(root, text="Сформировать документ", command=generate_document)
generate_button.pack()
generate_button.place(x=117, y=650)


def set_light_theme():
    customtkinter.set_appearance_mode("Dark")
    dark_on = customtkinter.CTkButton(root, width=35, height=35, text="", command=set_dark_theme, image=image_2)
    dark_on.pack()
    dark_on.place(x=340, y=700)


def set_dark_theme():
    customtkinter.set_appearance_mode("light")
    dark_off = customtkinter.CTkButton(root, width=35, height=35, text="", command=set_light_theme, image=image_1)
    dark_off.pack()
    dark_off.place(x=340, y=700)


dark_on = customtkinter.CTkButton(root, width=35, height=35, text="", command=set_dark_theme, image=image_2)
dark_on.pack()
dark_on.place(x=340, y=700)


# Второе диалоговое окно
def change():
    if type_var.get() == "Служебная записка":
        doc = Document("СЗ_шаблон.docx")
    elif type_var.get() == "Приказ":
        doc = Document("Приказ_шаблон.docx")
    elif type_var.get() == "Письмо":
        doc = Document("Письмо_шаблон.docx")

    def change_isp():
        section = doc.sections[-1]
        footer = section.footer
        footer_para = footer.paragraphs[-1]
        footer_para.text = "Исп. " + change_fio_entry.get() + ' ' + 'тел. ' + phone_entry.get()
        run = footer_para.runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(10)
        doc.save('Temp.docx')
        change_main.destroy()
        CTkMessagebox.CTkMessagebox(title="Docs Generator", message="Исполнитель успешно задан!")

    change_main = customtkinter.CTkToplevel()
    change_main.geometry("300x170+{}+{}".format(int(change_main.winfo_screenwidth() / 2 - 150),
                                                int(change_main.winfo_screenheight() / 2 - 100)))
    change_main.resizable(width=False, height=False)
    change_main.title("Указать исполнителя")

    change_fio = customtkinter.CTkLabel(change_main, width=100, height=30, text="ФИО исполнителя")
    change_fio.pack()

    change_fio_entry = customtkinter.CTkEntry(change_main, width=250, height=30)
    change_fio_entry.pack()

    phone = customtkinter.CTkLabel(change_main, width=100, height=30, text="Номер раб. телефона")
    phone.pack()

    phone_entry = customtkinter.CTkEntry(change_main, width=250, height=30)
    phone_entry.pack()

    change_button = customtkinter.CTkButton(change_main, height=30, text="Сменить", command=change_isp)
    change_button.pack()
    change_button.place(x=85, y=130)


change = customtkinter.CTkButton(root, width=40, height=35, text="Указать исполнителя", command=change)
change.pack()
change.place(x=15, y=700)

root.mainloop()
